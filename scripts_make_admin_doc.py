from pathlib import Path
from docx import Document
from docx.shared import Pt

md_path = Path(r"f:\serene\SERENE_ADMIN_GUIDE.md")
out_path = Path(r"f:\serene\SERENE_ADMIN_GUIDE.docx")

lines = md_path.read_text(encoding="utf-8").splitlines()
doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

i = 0
while i < len(lines):
    line = lines[i].rstrip()
    if not line.strip():
        doc.add_paragraph("")
        i += 1
        continue

    if line.startswith("---"):
        i += 1
        continue

    if line.startswith("## "):
        text = line[3:].strip()
        # If it's the first heading after title, make as heading 1
        if text.lower().startswith("admin operations guide"):
            doc.add_heading(text, level=1)
        else:
            doc.add_heading(text, level=2)
        i += 1
        continue

    if line.startswith("# "):
        doc.add_heading(line[2:].strip(), level=1)
        i += 1
        continue

    if line.lstrip().startswith("- "):
        doc.add_paragraph(line.lstrip()[2:].strip(), style="List Bullet")
        i += 1
        continue

    # numbered lines like "1. ..."
    stripped = line.strip()
    if len(stripped) > 2 and stripped[0].isdigit() and stripped[1] == ".":
        doc.add_paragraph(stripped, style="List Number")
        i += 1
        continue

    doc.add_paragraph(line)
    i += 1

doc.save(out_path)
print(f"Created {out_path}")
